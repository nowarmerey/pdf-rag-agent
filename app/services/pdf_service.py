import fitz  # PyMuPDF
import docx
import google.generativeai as genai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import CHUNK_SIZE, CHUNK_OVERLAP, GEMINI_API_KEY
import base64
import os

genai.configure(api_key=GEMINI_API_KEY)
vision_model = genai.GenerativeModel('gemini-2.0-flash')

# ═══════════════════════════════════════════
# Text Extraction
# ═══════════════════════════════════════════

def extract_text_from_pdf(file_path: str) -> str:
    """استخراج النص من PDF"""
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    doc.close()
    return full_text

def extract_text_from_word(file_path: str) -> str:
    """استخراج النص من Word"""
    doc = docx.Document(file_path)
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"
    # استخراج النص من الجداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + " "
            full_text += "\n"
    return full_text

def extract_text_from_image(file_path: str) -> str:
    """استخراج النص من الصور باستخدام Gemini Vision"""
    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    extension = os.path.splitext(file_path)[1].lower()
    mime_types = {
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png":  "image/png",
        ".webp": "image/webp"
    }
    mime_type = mime_types.get(extension, "image/jpeg")

    response = vision_model.generate_content([
        {
            "inline_data": {
                "mime_type": mime_type,
                "data": image_data
            }
        },
        """Extract ALL text from this image completely and accurately.
        If this is a legal document, contract, or official letter:
        - Extract every word, number, date, and signature
        - Maintain the original structure and formatting
        - Include headers, paragraphs, tables, and footnotes
        - Preserve paragraph numbers and article references
        Output ONLY the extracted text, nothing else."""
    ])
    return response.text

def extract_text(file_path: str, filename: str) -> str:
    """استخراج النص حسب نوع الملف"""
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in [".docx", ".doc"]:
        return extract_text_from_word(file_path)
    elif extension in [".jpg", ".jpeg", ".png", ".webp"]:
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

# ═══════════════════════════════════════════
# Text Splitting
# ═══════════════════════════════════════════

def split_text_into_chunks(text: str, filename: str) -> list[dict]:
    """تقسيم النص إلى chunks"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " "]
    )
    chunks = splitter.split_text(text)

    return [
        {
            "text": chunk,
            "metadata": {
                "source":      filename,
                "chunk_index": i
            }
        }
        for i, chunk in enumerate(chunks)
    ]